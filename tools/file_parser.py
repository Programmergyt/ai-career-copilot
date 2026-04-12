"""文件解析工具 — 支持 PDF / DOCX / Markdown / TXT。"""

from __future__ import annotations

from pathlib import Path

from log import get_logger

logger = get_logger("app")


def parse_file(file_path: str | Path) -> str:
    """解析文件内容为纯文本/Markdown。"""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(path)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(path)
    elif suffix in (".md", ".markdown", ".txt", ".text"):
        return _parse_text(path)
    else:
        logger.warning("Unsupported file type: %s, treating as plain text", suffix)
        return _parse_text(path)


def parse_content(content: str, filename: str = "") -> str:
    """解析内容字符串（用于 API 上传的 base64 解码后内容）。"""
    if filename:
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            return _parse_pdf_bytes(content.encode("latin-1"))
        elif suffix in (".docx", ".doc"):
            return _parse_docx_bytes(content.encode("latin-1"))
    return content


def _parse_pdf(path: Path) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        result = "\n\n".join(text_parts)
        logger.info("Parsed PDF: %s (%d chars)", path.name, len(result))
        return result
    except ImportError:
        logger.error("pdfplumber not installed")
        raise
    except Exception as e:
        logger.error("Failed to parse PDF %s: %s", path, e)
        raise


def _parse_pdf_bytes(data: bytes) -> str:
    import io
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
    return "\n\n".join(text_parts)


def _parse_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
        result = "\n\n".join(text_parts)
        logger.info("Parsed DOCX: %s (%d chars)", path.name, len(result))
        return result
    except ImportError:
        logger.error("python-docx not installed")
        raise
    except Exception as e:
        logger.error("Failed to parse DOCX %s: %s", path, e)
        raise


def _parse_docx_bytes(data: bytes) -> str:
    import io
    from docx import Document
    doc = Document(io.BytesIO(data))
    text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(text_parts)


def _parse_text(path: Path) -> str:
    text = path.read_text(encoding="utf-8")
    logger.info("Parsed text: %s (%d chars)", path.name, len(text))
    return text
